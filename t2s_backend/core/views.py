import io
import json
import uuid
from pathlib import Path

from asgiref.sync import async_to_sync
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Inches, RGBColor
from drf_spectacular.types import OpenApiTypes
from drf_spectacular.utils import OpenApiParameter, extend_schema, extend_schema_view
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from rest_framework import permissions, status, viewsets
from rest_framework.decorators import action
from rest_framework.parsers import FormParser, JSONParser, MultiPartParser
from rest_framework.response import Response
from rest_framework.views import APIView

from core.models import Chat, DatabaseConnection, Role, Message
from core.serializers import (
    ChatSerializer,
    DatabaseConnectionResponseSerializer,
    DatabaseConnectionSerializer,
    MessageDetailResponseSerializer,
    MessagePreviewSerializer,
)
from core.utils.db_connection import build_async_database_url, check_connection, set_active_connection_payload
from core.utils.fetch_data import fetch_data

PDF_FONT_NAME = "DejaVuSans"
PDF_FONT_BOLD_NAME = "DejaVuSans-Bold"
PDF_FONT_CANDIDATES = [
    (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ),
]


@extend_schema(tags=["Чаты"])
@extend_schema_view(
    list=extend_schema(summary="Получить список всех чатов"),
    create=extend_schema(summary="Создать новый чат"),
    retrieve=extend_schema(summary="Получить данные конкретного чата"),
    update=extend_schema(summary="Полное обновление чата"),
    partial_update=extend_schema(summary="Частичное изменение чата"),
    destroy=extend_schema(summary="Удалить чат")
)
class ChatViewSet(viewsets.ModelViewSet):
    queryset = Chat.objects.all()
    serializer_class = ChatSerializer
    permission_classes = [permissions.AllowAny]

    @extend_schema(
        summary="История сообщений пользователя",
        description="Возвращает только сообщения с ролью USER для конкретного чата",
        responses={200: MessagePreviewSerializer(many=True)},
        tags=["Сообщения"],
    )
    @action(detail=True, methods=["get"], url_path="history")
    def history(self, request, pk=None):
        chat = self.get_object()
        messages = Message.objects.filter(chat=chat, role=Role.USER).order_by("-created_at")

        serializer = MessagePreviewSerializer(messages, many=True)
        return Response(serializer.data)


class DatabaseConnectionView(APIView):
    permission_classes = [permissions.AllowAny]
    parser_classes = [MultiPartParser, JSONParser]

    @extend_schema(
        summary="Подключить внешнюю БД",
        description=(
                "Создает и активирует подключение к PostgreSQL/MySQL/SQLite. "
                "Для SQLite передается файл базы данных"
        ),
        request=DatabaseConnectionSerializer,
        responses={201: DatabaseConnectionResponseSerializer},
        tags=["Подключения к БД"],
    )
    def post(self, request):
        serializer = DatabaseConnectionSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        instance = serializer.save(is_active=False)

        sqlite_file_name = instance.sqlite_file.name if instance.sqlite_file else None
        try:
            db_url = build_async_database_url(
                db_type=instance.db_type,
                username=instance.username,
                password=instance.password,
                database_name=instance.database_name,
                host=instance.host,
                port=instance.port,
                sqlite_file_path=instance.sqlite_file.path if instance.sqlite_file else "",
            )
            async_to_sync(check_connection)(db_url)
        except Exception as exc:
            if instance.sqlite_file:
                instance.sqlite_file.delete(save=False)
            instance.delete()
            return Response(
                {"error": f"Не удалось подключиться к базе данных: {str(exc)}"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        DatabaseConnection.objects.filter(is_active=True).exclude(id=instance.id).update(is_active=False)
        instance.is_active = True
        instance.save(update_fields=["is_active", "updated_at"])
        set_active_connection_payload(
            {
                "db_type": instance.db_type,
                "username": instance.username,
                "password": instance.password,
                "database_name": instance.database_name,
                "host": instance.host,
                "port": instance.port,
                "sqlite_file_path": instance.sqlite_file.path if instance.sqlite_file else "",
            }
        )

        response_data = DatabaseConnectionResponseSerializer(instance).data
        if sqlite_file_name:
            response_data["sqlite_file"] = sqlite_file_name

        return Response(response_data, status=status.HTTP_201_CREATED)

    @extend_schema(
        summary="Получить список всех подключений",
        responses={200: DatabaseConnectionResponseSerializer(many=True)},
        tags=["Подключения к БД"],
    )
    def get(self, request):
        connections = DatabaseConnection.objects.all().order_by('-created_at')
        serializer = DatabaseConnectionResponseSerializer(connections, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


class MessageDetailView(APIView):
    @extend_schema(
        summary="Детальная информация о запросе",
        description="Ищет сообщение по ID и подтягивает данные из внешней",
        responses={200: MessageDetailResponseSerializer},
        tags=["Сообщения"],
    )
    def get(self, request, message_id):
        instance = get_object_or_404(Message, message_id=message_id)
        sql_text = instance.message
        request_message = Message.objects.get(id=message_id)
        if not sql_text:
            return Response(
                {"error": "SQL for this message has not been generated yet"},
                status=status.HTTP_404_NOT_FOUND,
            )

        try:
            external_body = async_to_sync(fetch_data)(sql_text)
        except Exception as exc:
            return Response({"error": str(exc)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        serializer = MessagePreviewSerializer(instance)
        response_data = serializer.data
        response_data["payload"] = external_body
        response_data["request"] = request_message.message

        return Response(response_data)


class MessageExportView(APIView):
    @extend_schema(
        summary="Экспорт сообщений пользователя",
        description="Экспортирует данные в нужный формат (JSON, XLSX, DOCX, PDF)",
        responses={
            200: MessagePreviewSerializer,
            (200, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"): OpenApiTypes.BINARY,
            (200, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"): OpenApiTypes.BINARY,
            (200, "application/pdf"): OpenApiTypes.BINARY,
        },
        tags=["Сообщения"],
        parameters=[
            OpenApiParameter(
                name="fmt",
                type=OpenApiTypes.STR,
                location="query",
                description="Export file format",
                required=False,
                enum=["xlsx", "docx", "pdf"],
            ),
        ],
    )
    def get(self, request, message_id):
        instance = get_object_or_404(Message, message_id=message_id)
        sql_text = instance.message
        if not sql_text:
            return Response(
                {"error": "SQL for this message has not been generated yet"},
                status=status.HTTP_404_NOT_FOUND,
            )

        try:
            external_data = async_to_sync(fetch_data)(sql_text)
        except Exception as exc:
            return Response(
                {"error": f"SQL Error: {str(exc)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        serializer = MessagePreviewSerializer(instance)
        response_data = serializer.data
        response_data["payload"] = external_data

        export_format = request.query_params.get("fmt")
        if export_format:
            return self.handle_export(export_format, response_data)

        return Response(response_data)

    def handle_export(self, fmt, data):
        filename = f"message_{uuid.uuid4().hex}"
        if fmt == "xlsx":
            return self.export_excel(data, filename)
        if fmt == "docx":
            return self.export_docx(data, filename)
        if fmt == "pdf":
            return self.export_pdf(data, filename)
        return Response({"error": "Unsupported format"}, status=400)

    def _collect_payload_headers(self, payload):
        headers = []
        for row in payload:
            if not isinstance(row, dict):
                continue
            for key in row.keys():
                header = str(key)
                if header not in headers:
                    headers.append(header)
        return headers

    def _stringify_cell(self, value):
        if value is None:
            return ""
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False, indent=2)
        return str(value)

    def _prepare_payload_table(self, payload):
        if not isinstance(payload, list) or not payload:
            return [], []

        headers = self._collect_payload_headers(payload)
        if not headers:
            return ["value"], [[self._stringify_cell(item)] for item in payload]

        rows = []
        for item in payload:
            if isinstance(item, dict):
                rows.append([self._stringify_cell(item.get(header)) for header in headers])
            else:
                rows.append([self._stringify_cell(item)])
        return headers, rows

    def _build_export_context(self, data):
        headers, rows = self._prepare_payload_table(data.get("payload"))
        return {
            "id": data["id"],
            "sql": data["message"],
            "created_at": str(data["created_at"]),
            "description": data.get("description") or "",
            "payload_count": len(data["payload"]) if isinstance(data.get("payload"), list) else 0,
            "headers": headers,
            "rows": rows,
        }

    def _register_pdf_font(self):
        regular_font_path = None
        bold_font_path = None
        for regular_candidate, bold_candidate in PDF_FONT_CANDIDATES:
            if regular_candidate.exists() and bold_candidate.exists():
                regular_font_path = regular_candidate
                bold_font_path = bold_candidate
                break

        if regular_font_path is None or bold_font_path is None:
            raise FileNotFoundError("No Cyrillic-compatible TTF font found for PDF export")

        registered_fonts = pdfmetrics.getRegisteredFontNames()
        if PDF_FONT_NAME not in registered_fonts:
            pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, regular_font_path))
        if PDF_FONT_BOLD_NAME not in registered_fonts:
            pdfmetrics.registerFont(TTFont(PDF_FONT_BOLD_NAME, bold_font_path))

    def export_excel(self, data, filename):
        context = self._build_export_context(data)
        wb = Workbook()

        summary_ws = wb.active
        summary_ws.title = "Сводка"
        summary_ws.append(["Поле", "Значение"])
        summary_ws.append(["ID сообщения", context["id"]])
        summary_ws.append(["Дата создания", context["created_at"]])
        summary_ws.append(["Количество строк", context["payload_count"]])
        summary_ws.append(["Описание", context["description"]])
        summary_ws.append(["SQL-запрос", context["sql"]])

        for cell in summary_ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(fill_type="solid", fgColor="006B62")
        summary_ws.column_dimensions["A"].width = 18
        summary_ws.column_dimensions["B"].width = 110
        for row in summary_ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        payload_ws = wb.create_sheet("Данные")
        if context["headers"]:
            payload_ws.append(context["headers"])
            for cell in payload_ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(fill_type="solid", fgColor="006B62")
                cell.alignment = Alignment(wrap_text=True, vertical="top")
            for row in context["rows"]:
                payload_ws.append(row)
            payload_ws.freeze_panes = "A2"
            for column in payload_ws.columns:
                letter = column[0].column_letter
                max_length = max(
                    len(self._stringify_cell(cell.value))
                    for cell in column
                    if cell.value is not None
                )
                payload_ws.column_dimensions[letter].width = min(max(max_length + 2, 14), 40)
            for row in payload_ws.iter_rows(min_row=2):
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
        else:
            payload_ws.append(["Нет данных"])

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{filename}.xlsx"'},
        )

    def export_docx(self, data, filename):
        context = self._build_export_context(data)
        doc = Document()
        heading0 = doc.add_heading("Экспорт сообщения", 0)
        for run in heading0.runs:
            run.font.color.rgb = RGBColor(0x00, 0x6B, 0x62)
        doc.add_paragraph(f"ID сообщения: {context['id']}")
        doc.add_paragraph(f"Дата создания: {context['created_at']}")
        doc.add_paragraph(f"Количество строк: {context['payload_count']}")
        if context["description"]:
            doc.add_paragraph(f"Описание: {context['description']}")
        doc.add_paragraph(f"SQL-запрос: {context['sql']}")

        heading1 = doc.add_heading("Данные", level=1)
        for run in heading1.runs:
            run.font.color.rgb = RGBColor(0x00, 0x6B, 0x62)

        if context["headers"]:
            table = doc.add_table(rows=1, cols=len(context["headers"]))
            table.style = "Table Grid"
            for index, header in enumerate(context["headers"]):
                table.rows[0].cells[index].text = header
            for row in context["rows"]:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value
        else:
            doc.add_paragraph("Нет данных.")

        for section in doc.sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
        )

    def export_pdf(self, data, filename):
        context = self._build_export_context(data)
        self._register_pdf_font()
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36,
        )
        styles = getSampleStyleSheet()
        styles["Title"].fontName = PDF_FONT_NAME
        styles["BodyText"].fontName = PDF_FONT_NAME
        styles["Title"].leading = 18
        styles["BodyText"].leading = 12
        elements = [
            Paragraph("Экспорт сообщения", styles["Title"]),
            Spacer(1, 12),
            Paragraph(f"ID сообщения: {context['id']}", styles["BodyText"]),
            Paragraph(f"Дата создания: {context['created_at']}", styles["BodyText"]),
            Paragraph(f"Количество строк: {context['payload_count']}", styles["BodyText"]),
        ]
        if context["description"]:
            elements.append(Paragraph(f"Описание: {context['description']}", styles["BodyText"]))
        elements.extend(
            [
                Spacer(1, 8),
                Paragraph(f"SQL-запрос: {context['sql']}", styles["BodyText"]),
                Spacer(1, 12),
            ]
        )

        if context["headers"]:
            table_data = [context["headers"], *context["rows"]]
            payload_table = Table(table_data, repeatRows=1)
            payload_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#006B62")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), PDF_FONT_BOLD_NAME),
                        ("FONTNAME", (0, 1), (-1, -1), PDF_FONT_NAME),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ]
                )
            )
            elements.append(payload_table)
        else:
            elements.append(Paragraph("Нет данных.", styles["BodyText"]))

        doc.build(elements)
        buffer.seek(0)
        return HttpResponse(
            buffer,
            content_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'},
        )

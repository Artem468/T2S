from io import BytesIO
from pathlib import Path
from unittest.mock import patch

import reportlab
from django.test import SimpleTestCase
from docx import Document
from openpyxl import load_workbook

from core import views
from core.views import MessageExportView


class MessageExportViewTests(SimpleTestCase):
    def setUp(self):
        self.view = MessageExportView()
        self.data = {
            "id": 42,
            "message": "select id, name, extra from users",
            "description": "Экспорт пользователей",
            "created_at": "2026-04-23T17:00:00",
            "payload": [
                {"id": 1, "name": "Alice", "extra": {"city": "Moscow"}},
                {"id": 2, "name": "Bob", "age": 30},
            ],
        }

    def test_prepare_payload_table_merges_all_payload_keys(self):
        headers, rows = self.view._prepare_payload_table(self.data["payload"])

        self.assertEqual(headers, ["id", "name", "extra", "age"])
        self.assertEqual(rows[0][0:2], ["1", "Alice"])
        self.assertIn('"city": "Moscow"', rows[0][2])
        self.assertEqual(rows[1][3], "30")

    def test_export_excel_creates_russian_named_sheets(self):
        response = self.view.export_excel(self.data, "report")
        workbook = load_workbook(BytesIO(response.content))

        self.assertEqual(response["Content-Type"], "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.assertEqual(workbook.sheetnames, ["Сводка", "Данные"])
        self.assertEqual(workbook["Сводка"]["A1"].value, "Поле")
        self.assertEqual(workbook["Сводка"]["A2"].value, "ID сообщения")
        self.assertEqual(workbook["Данные"]["A1"].value, "id")
        self.assertEqual(workbook["Данные"]["B2"].value, "Alice")

    def test_export_docx_contains_russian_headings(self):
        response = self.view.export_docx(self.data, "report")
        document = Document(BytesIO(response.content))

        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(document.paragraphs[0].text, "Экспорт сообщения")
        self.assertEqual(document.tables[0].cell(0, 0).text, "id")
        self.assertEqual(document.tables[0].cell(1, 1).text, "Alice")

    def test_export_pdf_returns_binary_pdf(self):
        font_dir = Path(reportlab.__file__).resolve().parent / "fonts"
        font_candidates = [(font_dir / "Vera.ttf", font_dir / "VeraBd.ttf")]
        with patch.object(views, "PDF_FONT_CANDIDATES", font_candidates), patch.object(
            views, "PDF_FONT_NAME", "TestVera"
        ), patch.object(views, "PDF_FONT_BOLD_NAME", "TestVeraBd"):
            response = self.view.export_pdf(self.data, "report")

        self.assertEqual(response["Content-Type"], "application/pdf")
        self.assertTrue(response.content.startswith(b"%PDF"))
